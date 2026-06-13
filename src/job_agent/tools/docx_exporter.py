"""DOCX exporter: cover letter and CV variant export using python-docx."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from job_agent.models.profile import Profile


class DocxExportError(Exception):
    """Raised when a DOCX export operation fails."""


def _ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def export_cover_letter(text: str, out_path: Path) -> Path:
    """Write a plain-text cover letter to a .docx file.

    One paragraph per line. Empty lines deliberately become blank paragraphs
    so the rendered Jinja template's section spacing (e.g. between Subject /
    opening / body / closing) carries into Word visually.
    Returns the absolute path. Ensures parent dir exists.
    """
    try:
        _ensure_parent(out_path)
        doc = Document()
        # Empty lines → blank paragraphs intentional (section spacing).
        for line in text.split("\n"):
            doc.add_paragraph(line)
        doc.save(str(out_path))
    except DocxExportError:
        raise
    except Exception as exc:
        raise DocxExportError(f"Failed to export cover letter: {exc}") from exc
    return out_path.resolve()


def export_cv_variant(profile: Profile, out_path: Path) -> Path:
    """Render a Profile to a structured .docx CV.

    Layout:
      H1: candidate name placeholder (or 'Curriculum Vitae')
      H2: Summary
      H2: Skills (highlighted_skills first, then the rest, deduped)
      H2: Experience
      H2: Education
      H2: Languages
    Returns the absolute path.
    """
    try:
        _ensure_parent(out_path)
        doc = Document()

        # H1 — title
        doc.add_heading("Curriculum Vitae", level=1)

        # Summary
        doc.add_heading("Summary", level=2)
        doc.add_paragraph(profile.summary or "")

        # Skills — highlighted first, then remaining (deduped, original order)
        doc.add_heading("Skills", level=2)
        ordered_skills: list[str] = list(profile.highlighted_skills)
        seen: set[str] = set(profile.highlighted_skills)
        for skill in profile.skills:
            if skill not in seen:
                ordered_skills.append(skill)
                seen.add(skill)
        for skill in ordered_skills:
            doc.add_paragraph(skill, style="List Bullet")

        # Experience
        doc.add_heading("Experience", level=2)
        for exp in profile.experience:
            date_range = f"{exp.start or '?'} – {exp.end or 'present'}"
            heading_text = f"{exp.title} at {exp.company}, {date_range}"
            p = doc.add_paragraph()
            run = p.add_run(heading_text)
            run.bold = True
            if exp.description:
                doc.add_paragraph(exp.description)

        # Education
        doc.add_heading("Education", level=2)
        for edu in profile.education:
            parts = [edu.degree]
            if edu.field:
                parts.append(f"in {edu.field}")
            parts.append(f"at {edu.institution}")
            if edu.start or edu.end:
                parts.append(f"({edu.start or '?'} – {edu.end or 'present'})")
            doc.add_paragraph(" ".join(parts))

        # Languages
        doc.add_heading("Languages", level=2)
        for lang in profile.languages:
            doc.add_paragraph(lang, style="List Bullet")

        doc.save(str(out_path))
    except DocxExportError:
        raise
    except Exception as exc:
        raise DocxExportError(f"Failed to export CV variant: {exc}") from exc

    return out_path.resolve()


_SLUG_STRIP_RE = re.compile(r"[^a-z0-9]+")


def make_slug(text: str, max_len: int = 40, fallback: str = "untitled") -> str:
    """Convert *text* to a lowercase alnum-and-dash slug, max *max_len* chars.

    Empty or whitespace-only input (or input that strips to nothing — e.g. only
    umlauts that get collapsed to dashes and then stripped) returns *fallback*
    so callers cannot silently produce filenames with empty slug segments.
    """
    lowered = (text or "").lower()
    slugged = _SLUG_STRIP_RE.sub("-", lowered).strip("-")
    slugged = slugged[:max_len].rstrip("-")
    return slugged or fallback
