"""Tests for job_agent.tools.docx_exporter."""

from __future__ import annotations

from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document

from job_agent.models.profile import Education, Experience, Profile
from job_agent.tools.docx_exporter import (
    DocxExportError,
    export_cover_letter,
    export_cv_variant,
    make_slug,
)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_full_profile() -> Profile:
    return Profile(
        summary="Experienced Python developer with 5 years of backend work.",
        skills=["Python", "SQL", "Docker", "REST APIs", "Git"],
        highlighted_skills=["Python", "SQL"],
        experience=[
            Experience(
                title="Backend Engineer",
                company="TechCorp",
                start="2020-01",
                end="2024-01",
                description="Built scalable REST APIs.",
            )
        ],
        education=[
            Education(
                degree="B.Sc. Computer Science",
                institution="Example University",
                field="Computer Science",
                start="2016",
                end="2020",
            )
        ],
        languages=["English", "German"],
    )


# ---------------------------------------------------------------------------
# Cover letter tests
# ---------------------------------------------------------------------------

def test_export_cover_letter_writes_file(tmp_path: Path) -> None:
    out = tmp_path / "cl.docx"
    result = export_cover_letter("hello\nworld", out)

    assert result.exists()
    doc = Document(result)
    texts = [p.text for p in doc.paragraphs]
    assert "hello" in texts
    assert "world" in texts


def test_export_cover_letter_paragraph_count(tmp_path: Path) -> None:
    out = tmp_path / "cl2.docx"
    export_cover_letter("hello\nworld", out)
    doc = Document(out)
    # Exactly 2 paragraphs for 2 lines
    assert len(doc.paragraphs) == 2
    assert doc.paragraphs[0].text == "hello"
    assert doc.paragraphs[1].text == "world"


# ---------------------------------------------------------------------------
# CV variant tests
# ---------------------------------------------------------------------------

def test_export_cv_variant_writes_file(tmp_path: Path) -> None:
    profile = _make_full_profile()
    out = tmp_path / "cv.docx"
    result = export_cv_variant(profile, out)

    assert result.exists()
    doc = Document(result)
    all_text = " ".join(p.text for p in doc.paragraphs)

    # All section headings present
    assert "Summary" in all_text
    assert "Skills" in all_text
    assert "Experience" in all_text
    assert "Education" in all_text
    assert "Languages" in all_text

    # Content present
    assert "Experienced Python developer" in all_text
    assert "Backend Engineer" in all_text
    assert "B.Sc. Computer Science" in all_text
    assert "English" in all_text


def test_export_cv_variant_highlighted_skills_first(tmp_path: Path) -> None:
    """highlighted_skills must appear before non-highlighted skills."""
    profile = _make_full_profile()
    # highlighted: Python, SQL — non-highlighted: Docker, REST APIs, Git
    out = tmp_path / "cv_order.docx"
    export_cv_variant(profile, out)

    doc = Document(out)
    # Find index of Skills heading paragraph, then collect bullet paragraphs
    paragraphs = doc.paragraphs
    skills_idx = next(i for i, p in enumerate(paragraphs) if p.text == "Skills")
    # Collect skill bullet texts until next heading
    skill_paras: list[str] = []
    for p in paragraphs[skills_idx + 1 :]:
        if p.style.name.startswith("Heading"):
            break
        if p.text:
            skill_paras.append(p.text)

    assert skill_paras[0] == "Python"
    assert skill_paras[1] == "SQL"
    # Non-highlighted come after
    assert "Docker" in skill_paras
    assert skill_paras.index("Python") < skill_paras.index("Docker")


def test_export_cv_variant_no_skill_duplication(tmp_path: Path) -> None:
    """Skills that appear in both highlighted_skills and skills must not be duplicated."""
    profile = Profile(
        summary="Dev",
        skills=["Python", "SQL", "Docker"],
        highlighted_skills=["Python", "SQL"],  # overlaps with skills
    )
    out = tmp_path / "cv_dedup.docx"
    export_cv_variant(profile, out)

    doc = Document(out)
    paragraphs = doc.paragraphs
    skills_idx = next(i for i, p in enumerate(paragraphs) if p.text == "Skills")
    skill_paras: list[str] = []
    for p in paragraphs[skills_idx + 1 :]:
        if p.style.name.startswith("Heading"):
            break
        if p.text:
            skill_paras.append(p.text)

    assert skill_paras.count("Python") == 1
    assert skill_paras.count("SQL") == 1


# ---------------------------------------------------------------------------
# Parent dir creation
# ---------------------------------------------------------------------------

def test_export_creates_parent_dir(tmp_path: Path) -> None:
    nested = tmp_path / "deep" / "nested" / "cl.docx"
    assert not nested.parent.exists()
    export_cover_letter("test", nested)
    assert nested.exists()


def test_export_cv_creates_parent_dir(tmp_path: Path) -> None:
    nested = tmp_path / "deep2" / "cv.docx"
    export_cv_variant(Profile(), nested)
    assert nested.exists()


# ---------------------------------------------------------------------------
# Failure wrapping
# ---------------------------------------------------------------------------

def test_make_slug_basic() -> None:
    assert make_slug("Hello World") == "hello-world"


def test_make_slug_lowercases() -> None:
    assert make_slug("ALLCAPS") == "allcaps"


def test_make_slug_strips_punctuation() -> None:
    assert make_slug("Senior Dev (m/w/d) — Berlin!") == "senior-dev-m-w-d-berlin"


def test_make_slug_handles_umlauts() -> None:
    """German umlauts must collapse to dashes (not crash, not leak through)."""
    out = make_slug("Bürokraft Köln (Voll/Teilzeit)")
    # Each ü/ö collapses to '-'; result is alnum+dash only.
    assert all(c.isalnum() or c == "-" for c in out)
    assert out.startswith("b-rokraft")


def test_make_slug_truncates_to_max_len() -> None:
    long_input = "a" * 100
    out = make_slug(long_input, max_len=40)
    assert len(out) <= 40


def test_make_slug_strips_trailing_dash_after_truncation() -> None:
    out = make_slug("python-developer-berlin", max_len=12)
    # 'python-devel' would end mid-word; ensure no trailing dash artifacts.
    assert not out.endswith("-")


def test_make_slug_empty_input_returns_fallback() -> None:
    assert make_slug("") == "untitled"


def test_make_slug_whitespace_only_returns_fallback() -> None:
    assert make_slug("   \t\n   ") == "untitled"


def test_make_slug_custom_fallback() -> None:
    assert make_slug("", fallback="job") == "job"


def test_make_slug_all_special_chars_returns_fallback() -> None:
    """Input that strips to nothing (e.g. only punctuation) returns fallback."""
    assert make_slug("!!!@@@---", fallback="company") == "company"


def test_export_failure_wrapped(tmp_path: Path) -> None:
    out = tmp_path / "fail.docx"
    boom = patch("job_agent.tools.docx_exporter.Document", side_effect=RuntimeError("boom"))
    with boom, pytest.raises(DocxExportError):
        export_cover_letter("text", out)
